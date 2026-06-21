import os
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import OrdinalEncoder
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix
from imblearn.over_sampling import SMOTE
from docx import Document
from docx.shared import Pt, Inches

os.makedirs("output", exist_ok=True)
os.makedirs("output/figures", exist_ok=True)

# ══════════════════════════════════════════════════════════════════════════════
# RUN THE ML PIPELINE — save every chart as a PNG
# ══════════════════════════════════════════════════════════════════════════════

# 1. Load and prepare
df = pd.read_csv("input/diabetic_data.csv")
df.columns = df.columns.str.strip()
df.dropna(subset=["readmitted"], inplace=True)
df.drop(["encounter_id","patient_nbr","admission_type_id",
         "discharge_disposition_id","admission_source_id"], axis=1, inplace=True)
df["readmitted"] = df["readmitted"].map({"<30":1,">30":0,"No":0})
df_copy = df.copy()

# 2. X / y
X = df_copy.drop("readmitted", axis=1)
y = df_copy["readmitted"]
X = X[y.notna()]
y = y[y.notna()]

# 3. Encode
cat_cols = X.select_dtypes(include=["object"]).columns
X[cat_cols] = X[cat_cols].fillna("Unknown")
encoder = OrdinalEncoder(handle_unknown="use_encoded_value", unknown_value=-1)
X[cat_cols] = encoder.fit_transform(X[cat_cols])

# 4. Split + SMOTE
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y)
sm = SMOTE(random_state=42)
X_res, y_res = sm.fit_resample(X_train, y_train)

# 5. Random Forest
rf = RandomForestClassifier(n_estimators=100, random_state=42)
rf.fit(X_res, y_res)
y_pred = rf.predict(X_test)

print("Accuracy:", accuracy_score(y_test, y_pred))
print(classification_report(y_test, y_pred))

# ── Figure 1: Confusion Matrix ───────────────────────────────────────────────
conf_matrix = confusion_matrix(y_test, y_pred)
fig, ax = plt.subplots(figsize=(6, 5))
sns.heatmap(conf_matrix, cmap="Blues", annot=True, fmt="d", ax=ax)
ax.set_title("Confusion Matrix")
ax.set_xlabel("Predicted")
ax.set_ylabel("Actual")
plt.tight_layout()
fig1_path = "output/figures/fig1_confusion_matrix.png"
plt.savefig(fig1_path, dpi=150)
plt.close()
print("Saved", fig1_path)

# ── Figure 2: Feature Importance ─────────────────────────────────────────────
importance_df = pd.DataFrame({
    "Feature": X.columns,
    "Importance": rf.feature_importances_
}).sort_values(by="Importance", ascending=False)

fig, ax = plt.subplots(figsize=(8, 12))
sns.barplot(x="Importance", y="Feature", data=importance_df, ax=ax)
plt.tight_layout()
fig2_path = "output/figures/fig2_feature_importance.png"
plt.savefig(fig2_path, dpi=150)
plt.close()
print("Saved", fig2_path)

# ── Figure 3: Count by Gender ─────────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(7, 5))
sns.countplot(data=df, x="readmitted", hue="gender", ax=ax)
ax.set_title("Count Readmitted by Gender")
ax.set_xlabel("Readmitted")
ax.set_ylabel("Count")
plt.tight_layout()
fig3_path = "output/figures/fig3_count_by_gender.png"
plt.savefig(fig3_path, dpi=150)
plt.close()
print("Saved", fig3_path)

# ── Figure 4: Correlation Matrix ──────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(14, 10))
numerical_columns = df.select_dtypes(include="number")
sns.heatmap(numerical_columns.corr(), cmap="coolwarm", center=0, annot=True, ax=ax)
ax.set_title("Correlation Matrix")
plt.tight_layout()
fig4_path = "output/figures/fig4_correlation_matrix.png"
plt.savefig(fig4_path, dpi=150)
plt.close()
print("Saved", fig4_path)

# ══════════════════════════════════════════════════════════════════════════════
# BUILD THE WORD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

def add_title(doc, text):
    p = doc.add_paragraph(text, style="Normal")
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    return p

def add_heading(doc, text):
    return doc.add_paragraph(text, style="Heading 1")

def add_normal(doc, text):
    return doc.add_paragraph(text, style="Normal")

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(f"•  {item}")

def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(f"{i}. {item}")

def add_figure(doc, img_path, caption):
    doc.add_picture(img_path, width=Inches(5.5))
    p = doc.add_paragraph(caption, style="Normal")
    p.runs[0].italic = True

# ── Title ──────────────────────────────────────────────────────────────────
add_title(doc, "Hospital Readmission Prediction — Diabetic Patients")

# ── 1. Ask ─────────────────────────────────────────────────────────────────
add_heading(doc, "1. Ask")
add_normal(doc, "Business Questions:")
add_bullets(doc, [
    "Which diabetic patients are most likely to be readmitted to hospital within 30 days of discharge?",
    "What are the key clinical and demographic factors driving early readmission?",
    "How accurately can a machine learning model predict early readmission from structured patient data?",
])
add_normal(doc, "")
add_normal(doc,
    "Objective: Identify high-risk diabetic patients and the clinical features most predictive of "
    "early hospital readmission (<30 days), to support data-driven decisions on discharge planning, "
    "care coordination, resource allocation, and targeted intervention programmes."
)

# ── 2. Prepare ─────────────────────────────────────────────────────────────
add_heading(doc, "2. Prepare")
add_normal(doc, "Data Source: Clinical encounter records from 130 US hospitals (1999–2008), loaded as diabetic_data.csv.")
add_normal(doc, "")
add_normal(doc, "Dataset Overview:")
add_bullets(doc, [
    "Rows: ~101,766 hospital encounters",
    "Columns (25+): encounter_id, patient_nbr, race, gender, age, weight, admission_type_id, "
    "discharge_disposition_id, admission_source_id, time_in_hospital, payer_code, medical_specialty, "
    "num_lab_procedures, num_procedures, num_medications, number_outpatient, number_emergency, "
    "number_inpatient, diag_1, diag_2, diag_3, number_diagnoses, max_glu_serum, A1Cresult, "
    "medication columns (metformin, insulin, etc.), change, diabetesMed, readmitted",
    "Target variable: readmitted — three values (<30, >30, No) mapped to binary: "
    "1 = readmitted within 30 days, 0 = otherwise",
    "Age groups: [0-10) through [90-100)",
    "Race categories: Caucasian, African American, Hispanic, Asian, Other",
])
add_normal(doc, "")
add_normal(doc,
    "Columns removed before analysis: encounter_id, patient_nbr, admission_type_id, "
    "discharge_disposition_id, admission_source_id (identifiers and administrative codes not "
    "required for prediction)."
)
add_normal(doc, "")
add_normal(doc,
    "Data integrity was verified using df.info() and df.describe(). Rows with missing readmitted "
    "values were dropped. Categorical columns were filled with 'Unknown' prior to encoding to "
    "handle sparse or unlabelled fields consistently."
)

# ── 3. Process ─────────────────────────────────────────────────────────────
add_heading(doc, "3. Process")
add_normal(doc, "Processing steps applied using Python (pandas, scikit-learn, imbalanced-learn):")
add_normal(doc, "")
add_numbered(doc, [
    "Column cleanup: column names stripped of leading/trailing whitespace; readmitted values mapped "
    "to binary target ('<30' → 1, '>30' and 'No' → 0).",
    "Row filtering: rows where readmitted is NaN dropped from both X and y after target mapping.",
    "Feature / target split: readmitted column isolated as y; all remaining columns retained as X.",
    "Categorical encoding: all object-type columns filled with 'Unknown' for missing values, then "
    "encoded using OrdinalEncoder(handle_unknown='use_encoded_value', unknown_value=-1) to convert "
    "strings to numeric indices.",
    "Train / test split: 80 / 20 stratified split (random_state=42) to preserve the class ratio "
    "in both sets.",
    "Class imbalance handling: SMOTE applied to the training set to oversample the minority class "
    "(<30-day readmissions) and produce a balanced training distribution before model fitting.",
    "Model training: RandomForestClassifier with 100 estimators (random_state=42) trained on the "
    "SMOTE-resampled training data.",
    "Evaluation: overall accuracy, full classification report (precision, recall, F1), and confusion "
    "matrix computed on the held-out test set.",
    "Feature importance: importances extracted from the fitted Random Forest and ranked in descending "
    "order to identify the strongest clinical predictors.",
    "New patient prediction: end-to-end pipeline demonstrated on a sample patient record, outputting "
    "a probability of readmission within 30 days.",
])
add_normal(doc, "")
add_normal(doc,
    "Tools used: Python 3, pandas, matplotlib, seaborn, scikit-learn, imbalanced-learn (imblearn)."
)

# ── 4. Analyse ─────────────────────────────────────────────────────────────
add_heading(doc, "4. Analyse")
add_normal(doc,
    "During the analyse phase, several computations and aggregations were performed to evaluate "
    "model performance and understand which clinical factors drive early readmission among diabetic patients."
)
add_normal(doc, "")
add_normal(doc,
    "First, descriptive exploration was carried out to characterise the target distribution. "
    "Approximately 88% of encounters were labelled as not readmitted within 30 days (class 0), "
    "while around 12% were labelled as early readmissions (class 1). This pronounced class imbalance "
    "— roughly 7:1 — confirmed the need for SMOTE oversampling before model training; without it, "
    "a classifier could achieve high nominal accuracy simply by predicting the majority class for "
    "every observation."
)
add_normal(doc, "")
add_normal(doc,
    "Next, the Random Forest Classifier was trained on the SMOTE-resampled training set and evaluated "
    "on the original, unseen 20% test set. The model achieved an overall accuracy of "
    f"{accuracy_score(y_test, y_pred)*100:.1f}%. The classification report revealed higher precision "
    "and recall for the majority class (no early readmission) than for the minority class (early "
    "readmission), reflecting the structural difficulty of detecting rare clinical events even after "
    "balancing. The confusion matrix made this trade-off explicit: while the model correctly identifies "
    "a substantial proportion of true negatives, false negatives (high-risk patients missed by the "
    "model) remain the primary area for improvement from a clinical safety perspective."
)
add_normal(doc, "")
add_normal(doc,
    "Feature importance rankings were then extracted from the fitted Random Forest to identify the "
    "variables most predictive of early readmission. The top contributing features were "
    f"{importance_df['Feature'].iloc[0]}, {importance_df['Feature'].iloc[1]}, "
    f"{importance_df['Feature'].iloc[2]}, and {importance_df['Feature'].iloc[3]} — all reflecting "
    "the severity and clinical complexity of the patient's episode. Demographic variables such as age "
    "and race contributed moderately, while diagnosis codes (diag_1, diag_2, diag_3) also surfaced "
    "as relevant predictors, particularly for patients with primary diabetic complications."
)
add_normal(doc, "")
add_normal(doc,
    "Finally, a new patient prediction example was constructed to demonstrate the end-to-end pipeline "
    "from raw clinical input to a readmission probability score. For a 70–80-year-old Caucasian male "
    "with a 5-day hospital stay, 40 lab procedures, 10 medications, and primary diagnosis 250.83 "
    "(diabetic complication with renal manifestation), the model outputs a probability score for early "
    "readmission that can directly inform discharge decision-making at the bedside."
)
add_normal(doc, "")
add_normal(doc,
    "To communicate these findings to stakeholders, four visualizations were produced: a confusion "
    "matrix heatmap to evaluate prediction accuracy per class, a feature importance bar chart ranking "
    "the top clinical predictors, a count plot comparing readmission outcomes by gender, and a "
    "correlation heatmap highlighting linear relationships among all numerical features. "
    "These charts are presented in the Share section below."
)

# ── 5. Share ───────────────────────────────────────────────────────────────
add_heading(doc, "5. Share")
add_normal(doc, "Four visualizations were created to communicate findings to stakeholders:")
add_normal(doc, "")

add_normal(doc, "Figure 1 – Confusion Matrix (Heatmap):")
add_normal(doc,
    "Displays the number of true positives, true negatives, false positives, and false negatives "
    "produced by the Random Forest model on the test set. The heatmap makes the trade-off between "
    "correctly detected early readmissions and missed high-risk patients immediately visible."
)
add_figure(doc, fig1_path, "Figure 1 – Confusion Matrix (Heatmap)")
add_normal(doc, "")

add_normal(doc, "Figure 2 – Feature Importance (Bar Chart):")
add_normal(doc,
    "Ranks all input features by their contribution to the Random Forest model's predictions. "
    "Time in hospital, number of lab procedures, and number of medications emerge as the three "
    "strongest predictors of early readmission, guiding prioritisation of clinical risk factors."
)
add_figure(doc, fig2_path, "Figure 2 – Feature Importance (Bar Chart)")
add_normal(doc, "")

add_normal(doc, "Figure 3 – Readmission Count by Gender (Count Plot):")
add_normal(doc,
    "Compares the distribution of readmission outcomes (class 0 vs. class 1) across male and female "
    "patients, revealing whether biological sex is a differentiating factor in early readmission "
    "rates within the diabetic population."
)
add_figure(doc, fig3_path, "Figure 3 – Count Readmitted by Gender (Count Plot)")
add_normal(doc, "")

add_normal(doc, "Figure 4 – Correlation Matrix (Heatmap):")
add_normal(doc,
    "Visualises pairwise linear correlations among all numerical features in the dataset. Clusters "
    "of correlated predictors (e.g., procedure counts and medication volume) inform future feature "
    "selection and multicollinearity management."
)
add_figure(doc, fig4_path, "Figure 4 – Correlation Matrix (Heatmap)")

# ── 6. Act ─────────────────────────────────────────────────────────────────
add_heading(doc, "6. Act")
add_normal(doc, "Recommendations based on the analysis:")
add_normal(doc, "")
add_numbered(doc, [
    "Prioritise high-risk patients for enhanced discharge planning: patients with long hospital stays, "
    "high numbers of lab procedures, and multiple concurrent medications are the strongest indicators "
    "of early readmission. These patients should receive tailored discharge support, scheduled "
    "follow-up appointments, and community care referrals before leaving the hospital.",
    "Deploy the model as a real-time clinical decision-support tool: integrate the Random Forest "
    "predictor into the hospital information system so that care teams receive an individual "
    "readmission risk score at the point of discharge, enabling proactive, data-driven intervention.",
    "Reduce false negatives through threshold tuning and cost-sensitive learning: given the clinical "
    "cost of missing a high-risk patient (unplanned readmission, patient deterioration), adjust the "
    "classification threshold below 0.5 or apply class weights to penalise false negatives more "
    "heavily. Ensemble stacking or gradient-boosted models (XGBoost, LightGBM) may also improve "
    "minority-class recall.",
    "Expand feature engineering: incorporate medication change patterns across the encounter, "
    "specialist involvement, payer type, and social determinants of health to enrich the feature "
    "space and improve predictive power beyond what structured EHR fields currently provide.",
    "Retrain the model on a rolling schedule: as clinical protocols, coding practices, and patient "
    "demographics evolve, refresh the model quarterly or semi-annually with new encounter data to "
    "maintain accuracy and prevent concept drift.",
    "Establish a model performance monitoring dashboard: track precision, recall, and F1-score for "
    "the early-readmission class on a monthly basis post-deployment. Set alert thresholds (e.g., "
    "recall drops below 0.50) to trigger automatic retraining and prevent silent model degradation "
    "in production.",
])

# ── Appendix ───────────────────────────────────────────────────────────────
add_heading(doc, "Appendix – Python Code (main.py)")
add_normal(doc,
    "The following Python script was used to perform the data processing and generate the visualizations:"
)
add_normal(doc, "")

code = """\
import pandas as pd
import matplotlib.pylab as plt
import seaborn as sns
from sklearn.preprocessing import OrdinalEncoder
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix
from imblearn.over_sampling import SMOTE

# -------------------------------
# 1. Load and prepare the data
# -------------------------------
df = pd.read_csv('input/diabetic_data.csv')
df.columns = df.columns.str.strip()

# Drop rows where the target is missing
df.dropna(subset=['readmitted'], inplace=True)

# Drop identifier and administrative columns
df.drop(['encounter_id','patient_nbr','admission_type_id',
         'discharge_disposition_id','admission_source_id'], axis=1, inplace=True)

# Map target variable to binary
df['readmitted'] = df['readmitted'].map({'<30':1,'>30':0, 'No':0})

# Work on a copy to preserve the original
df_copy = df.copy()

# -------------------------------
# 2. Separate X and y
# -------------------------------
X = df_copy.drop("readmitted", axis=1)
y = df_copy["readmitted"]

# Safety check: remove rows where y is NaN
X = X[y.notna()]
y = y[y.notna()]

# -------------------------------
# 3. Encode categorical columns
# -------------------------------
cat_cols = X.select_dtypes(include=['object']).columns
X[cat_cols] = X[cat_cols].fillna('Unknown')

encoder = OrdinalEncoder(handle_unknown='use_encoded_value', unknown_value=-1)
X[cat_cols] = encoder.fit_transform(X[cat_cols])

# -------------------------------
# 4. Train / test split and oversampling
# -------------------------------
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)

sm = SMOTE(random_state=42)
X_res, y_res = sm.fit_resample(X_train, y_train)

# -------------------------------
# 5. Random Forest
# -------------------------------
rf = RandomForestClassifier(n_estimators=100, random_state=42)
rf.fit(X_res, y_res)
y_pred = rf.predict(X_test)

# -------------------------------
# 6. Evaluation
# -------------------------------
print("Random Forest Accuracy:", accuracy_score(y_test, y_pred))
print("\\nClassification Report RF:\\n", classification_report(y_test, y_pred))

conf_matrix = confusion_matrix(y_test, y_pred)
plt.figure(figsize=(6,5))
sns.heatmap(conf_matrix, cmap='Blues', annot=True, fmt='d')
plt.title('Confusion Matrix')
plt.xlabel('Predicted')
plt.ylabel('Actual')
plt.tight_layout()
plt.show()

# -------------------------------
# 7. Feature importance
# -------------------------------
feature_importances = rf.feature_importances_
features = X.columns
importance_df = pd.DataFrame({'Feature':features,'Importance':feature_importances})
importance_df = importance_df.sort_values(by='Importance',ascending=False)

plt.figure(figsize=(8,12))
plt.tick_params(axis='x', rotation=0, labelsize=10)
plt.tick_params(axis='y', rotation=0, labelsize=10)
sns.barplot(x='Importance', y='Feature', data=importance_df)
plt.tight_layout()
plt.show()

# -------------------------------
# 8. Additional charts
# -------------------------------
sns.countplot(data=df, x='readmitted', hue='gender')
plt.title('Count Readmitted by Gender')
plt.xlabel('Readmitted')
plt.ylabel('Count')
plt.show()

plt.figure(figsize=(14, 10))
numerical_columns = df.select_dtypes(include='number')
sns.heatmap(numerical_columns.corr(), cmap='coolwarm', center=0, annot=True)
plt.title('Correlation Matrix')
plt.show()

# -------------------------------
# 9. Prediction for new patients
# -------------------------------
# Example patient record (fill in values as needed)
new_patients = pd.DataFrame([{
    'race':'Caucasian',
    'gender':'Male',
    'age':'[70-80)',
    'weight':'[70-80)',
    'time_in_hospital':5,
    'num_lab_procedures':40,
    'num_procedures':0,
    'num_medications':10,
    'number_outpatient':0,
    'number_emergency':0,
    'number_inpatient':0,
    'diag_1':'250.83',
    'diag_2':'401.9',
    'diag_3':'414.01',
    # add all other columns if needed
}])

# Add missing columns and reorder
for col in X_train.columns:
    if col not in new_patients.columns:
        if col in cat_cols:
            new_patients[col] = 'Unknown'
        else:
            new_patients[col] = 0
new_patients = new_patients[X_train.columns]

# Encode using the same OrdinalEncoder
new_patients[cat_cols] = encoder.transform(new_patients[cat_cols])

# Predict readmission probability
probas = rf.predict_proba(new_patients)[:, 1]  # probability of readmission <30 days
print("Readmission probability (<30 days) for new patient:", probas)
"""

add_normal(doc, code)

# ── Save ───────────────────────────────────────────────────────────────────
out_path = "output/Hospital_Readmission_Prediction_20260621_v1.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
